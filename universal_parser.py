import os
import json
import fitz  # PyMuPDF
import textract
import pandas as pd
import pdfplumber
import chardet
import subprocess
import numpy as np
import io
import threading
import pytesseract
from pdf2image import convert_from_path
import tempfile
import shutil
import subprocess
from scipy.stats import entropy
from PIL import Image
from docx.table import _Cell, Table
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from typing import Dict, Union, List
from PIL import Image
from concurrent.futures import ThreadPoolExecutor
from hashlib import md5
from docx import Document
from docx.text.paragraph import Paragraph as DocumentParagraph


def iter_block_items(parent):
    """
    Yield paragraphs and tables in document order from a parent element.
    """
    for child in parent.element.body.iterchildren():
        if child.tag.endswith('}p'):
            yield DocumentParagraph(parent._parent, child)
        elif child.tag.endswith('}tbl'):
            yield Table(child, parent)

def extract_visual_figures_from_docx(filepath, dpi=300, min_length=100, existing_hashes=None, save_dir="extracted_images/ocr_docx"):
    os.makedirs(save_dir, exist_ok=True)
    figures = []
    seen_hashes = existing_hashes or set()

    with tempfile.TemporaryDirectory() as tmpdir:
        # Convert to PDF first
        pdf_path = os.path.join(tmpdir, "temp.pdf")
        subprocess.run(["docx2pdf", filepath, pdf_path], check=True)
        
        # Convert PDF pages to images
        pages = convert_from_path(pdf_path, dpi=dpi)
        for page_num, img in enumerate(pages):
            buffer = io.BytesIO()
            img.save(buffer, format="PNG")
            img_bytes = buffer.getvalue()
            image_hash = get_image_hash(img_bytes)

            if image_hash in seen_hashes:
                continue
            seen_hashes.add(image_hash)

            img_path = os.path.join(save_dir, f"ocr_page{page_num+1}.png")
            img.save(img_path)

            ocr_text = ocr_image_with_tesseract(img)
            if ocr_text and len(ocr_text) >= min_length:
                figures.append({
                    "image": img_path,
                    "text": f"**Figure (Page {page_num+1})**:\n```\n{ocr_text}\n```"
                })
    return figures, seen_hashes

def extract_unique_images_from_docx(filepath, output_dir="extracted_images/docx"):
    os.makedirs(output_dir, exist_ok=True)
    doc = Document(filepath)
    images = []
    seen_hashes = set()

    for i, rel in enumerate(doc.part._rels):
        rel = doc.part._rels[rel]
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            image_hash = get_image_hash(image_data)
            if image_hash in seen_hashes or is_image_blank_or_small(image_data):
                continue
            seen_hashes.add(image_hash)
            image_ext = rel.target_ref.split('.')[-1]
            image_path = os.path.join(output_dir, f"image{i+1}.{image_ext}")
            with open(image_path, 'wb') as f:
                f.write(image_data)
            images.append(image_path)
    return images, seen_hashes


def ocr_image_with_tesseract(img: Image.Image) -> str:
    try:
        text = pytesseract.image_to_string(img)
        return text.strip()
    except Exception as e:
        print(f"OCR error: {e}")
        return ""

def extract_visual_figures_from_pdf(filepath, dpi=300, min_length=100, existing_hashes=None, save_dir="extracted_images/ocr"):
    os.makedirs(save_dir, exist_ok=True)
    figures = []
    seen_hashes = existing_hashes or set()
    doc = fitz.open(filepath)

    for page_num, page in enumerate(doc):
        # Extract text normally
        page_text = page.get_text().strip()

        # If page text is sufficient, skip OCR
        if page_text and len(page_text) >= min_length:
            continue  # Skip OCR image generation on this page

        # Otherwise, do OCR on page image
        pix = page.get_pixmap(dpi=dpi)
        img_bytes = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_bytes))
        image_hash = get_image_hash(img_bytes)

        if image_hash in seen_hashes:
            continue
        seen_hashes.add(image_hash)

        img_path = os.path.join(save_dir, f"ocr_page{page_num+1}.png")
        img.save(img_path)

        ocr_text = ocr_image_with_tesseract(img)
        if ocr_text and len(ocr_text) >= min_length:
            figures.append({
                "image": img_path,
                "text": f"**Figure (Page {page_num+1})**:\n```\n{ocr_text}\n```"
            })

    return figures, seen_hashes

def clean_text(text):
    return text.encode('utf-16', 'surrogatepass').decode('utf-16')

def detect_encoding(filepath):
    with open(filepath, 'rb') as f:
        result = chardet.detect(f.read())
    return result['encoding']

def parse_txt(filepath):
    encoding = detect_encoding(filepath)
    with open(filepath, 'r', encoding=encoding) as f:
        return {'text': f.read()}

def parse_csv(filepath):
    df = pd.read_csv(filepath)
    return {'text': df.to_string(index=False), 'tables': [df.to_dict(orient='records')]}

def is_image_blank(image_path, std_threshold=6, mean_threshold=245, entropy_threshold=1.0, resize_to=(64, 64)):
    try:
        img = Image.open(image_path).convert("L").resize(resize_to)
        arr = np.array(img)

        stddev = arr.std()
        mean = arr.mean()
        hist, _ = np.histogram(arr, bins=256, range=(0, 255), density=True)
        ent = entropy(hist + 1e-10)

        if stddev < std_threshold and mean > mean_threshold and ent < entropy_threshold:
            return True
        return False
    except:
        return False

def is_image_too_small(image_path, min_size=20):
    try:
        img = Image.open(image_path)
        width, height = img.size
        return width < min_size or height < min_size
    except:
        return False

def get_image_hash(image_bytes):
    return md5(image_bytes).hexdigest()

def extract_unique_images_from_pdf(filepath, output_dir="extracted_images/pdf"):
    os.makedirs(output_dir, exist_ok=True)
    doc = fitz.open(filepath)
    images_by_page = []
    seen_hashes = set()

    for page_num, page in enumerate(doc):
        page_images = []
        image_list = page.get_images(full=True)

        for img_index, img in enumerate(image_list):
            try:
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_hash = get_image_hash(image_bytes)

                if image_hash in seen_hashes or is_image_blank_or_small(image_bytes):
                    continue
                seen_hashes.add(image_hash)

                image_ext = base_image["ext"]
                image_filename = f"page{page_num+1}_img{img_index+1}.{image_ext}"
                image_path = os.path.join(output_dir, image_filename).replace("\\", "/")

                with open(image_path, "wb") as f:
                    f.write(image_bytes)

                page_images.append({
                    "path": image_path,
                    "page": page_num + 1
                })

            except Exception as e:
                print(f"Image extraction error on page {page_num+1}: {e}")

        images_by_page.extend(page_images)

    return images_by_page, seen_hashes

def process_image_from_pdf(doc, page_num, img_index, img, hashes, output_dir):
    try:
        xref = img[0]
        base_image = doc.extract_image(xref)
        image_bytes = base_image["image"]
        image_hash = get_image_hash(image_bytes)
        
        with hashes_lock:
            if image_hash in hashes:
                return None
            hashes.add(image_hash)
        
        if is_image_blank_or_small(image_bytes):
            return None

        image_ext = base_image["ext"]
        image_filename = f"page{page_num+1}_img{img_index+1}.{image_ext}"
        image_path = os.path.join(output_dir, image_filename)
        
        with open(image_path, "wb") as f:
            f.write(image_bytes)

        return {"path": image_path, "page": page_num + 1}
    
    except Exception as e:
        print(f"Error processing image on page {page_num+1} img {img_index+1}: {e}")
        return None


def is_image_blank_or_small(image_bytes, std_threshold=6, mean_threshold=245, entropy_threshold=1.0, resize_to=(64, 64), min_size=20):
    try:
        img = Image.open(io.BytesIO(image_bytes)).convert("L").resize(resize_to)
        arr = np.array(img)
        stddev = arr.std()
        mean = arr.mean()
        hist, _ = np.histogram(arr, bins=256, range=(0, 255), density=True)
        ent = entropy(hist + 1e-10)
        if stddev < std_threshold and mean > mean_threshold and ent < entropy_threshold:
            return True
        img = Image.open(io.BytesIO(image_bytes))
        width, height = img.size
        if width < min_size or height < min_size:
            return True
        return False
    except Exception as e:
        print(f"Image blank/small check error: {e}")
        return True  # treat errors as invalid image

hashes_lock = threading.Lock()


def extract_comments_from_docx(filepath):
    comments = []
    doc = Document(filepath)
    part = doc.part
    rels = part.rels

    for rel in rels:
        if rels[rel].reltype == RT.COMMENTS:
            comments_part = rels[rel].target_part
            comments_xml = comments_part._blob.decode('utf-8')
            import xml.etree.ElementTree as ET
            tree = ET.fromstring(comments_xml)
            for comment in tree.findall('.//w:comment', namespaces={'w': qn('w')}):
                text = ''.join(comment.itertext()).strip()
                if text:
                    comments.append(text)
    return comments

def parse_docx_ordered(filepath):
    content = []
    images, image_hashes = extract_unique_images_from_docx(filepath)
    visual_figures, _ = extract_visual_figures_from_docx(filepath, existing_hashes=image_hashes)

    doc = Document(filepath)
    # We will collect images separately with a mapping to placeholders if possible
    # Since inline image extraction is tricky, for now just list extracted images separately.

    for block in iter_block_items(doc):
        if isinstance(block, DocumentParagraph):
            text = block.text.strip()
            if text:
                content.append({"type": "text", "content": text})
        elif isinstance(block, Table):
            # Extract table content as list of rows
            table_data = []
            for row in block.rows:
                table_data.append([cell.text.strip() for cell in row.cells])
            content.append({"type": "table", "content": table_data})

    # Add extracted images as separate entries
    for img_path in images:
        content.append({"type": "image", "content": img_path})

    # Add visual figures (OCR) as text blocks
    for vf in visual_figures:
        content.append({"type": "visual_figure", "content": vf["text"]})

    # Extract comments
    annotations = extract_comments_from_docx(filepath)
    for note in annotations:
        content.append({"type": "annotation", "content": note})

    return content

def extract_annotations_from_pdf(filepath):
    annotations = []
    doc = fitz.open(filepath)
    for page in doc:
        annots = page.annots()
        if annots:
            for annot in annots:
                info = annot.info
                if 'content' in info:
                    annotations.append(info['content'])
    return annotations

def parse_pdf_ordered(filepath):
    content = []
    annotations = extract_annotations_from_pdf(filepath)
    images, image_hashes = extract_unique_images_from_pdf(filepath)
    visual_figures, _ = extract_visual_figures_from_pdf(filepath, existing_hashes=image_hashes)

    with pdfplumber.open(filepath) as pdf:
        for page_num, page in enumerate(pdf.pages):
            # Text
            text = page.extract_text()
            if text:
                content.append({"type": "text", "content": text.strip(), "page": page_num+1})

            # Tables
            try:
                tables = page.extract_tables()
                for table in tables:
                    content.append({"type": "table", "content": table, "page": page_num+1})
            except:
                pass

            # Images on this page
            # Images on this page
            page_images = [img["path"] for img in images if img["page"] == page_num + 1]
            for img_path in page_images:
                content.append({"type": "image", "content": img_path, "page": page_num + 1})

            # Annotations on this page
            # We can try to match annotations to pages, if available, else add at end
            # For now add all at the end after all pages

            # Visual figures OCR on this page
            # Visual figures may have page info, filter those for this page
            vf_for_page = [vf["text"] for vf in visual_figures if f"(Page {page_num+1})" in vf["text"]]
            for vf_text in vf_for_page:
                content.append({"type": "visual_figure", "content": vf_text, "page": page_num+1})

    # Append annotations at the end (could be improved by linking to page)
    for note in annotations:
        content.append({"type": "annotation", "content": note})

    return content


def parse_with_textract(filepath):
    try:
        text = textract.process(filepath).decode("utf-8")
        return {'text': text}
    except Exception as e:
        return {'error': f"Textract failed: {e}"}

def convert_with_libreoffice(filepath) -> Union[str, None]:
    try:
        output_dir = "/tmp"
        subprocess.run(["libreoffice", "--headless", "--convert-to", "docx", filepath, "--outdir", output_dir],
                       check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        filename = os.path.basename(filepath).rsplit('.', 1)[0] + '.docx'
        return os.path.join(output_dir, filename)
    except Exception as e:
        print(f"LibreOffice conversion failed: {e}")
        return None

def parse_xlsx(filepath: str) -> Dict[str, Union[str, List]]:
    xlsx_data = {
        "tables": [],
        "headings": [],
        "text": ""
    }

    try:
        xls = pd.read_excel(filepath, sheet_name=None, engine="openpyxl")
        xlsx_data["headings"] = list(xls.keys())

        for sheet_name, df in xls.items():
            if df.empty:
                continue
            df = df.fillna("").astype(str)
            table = [df.columns.tolist()] + df.values.tolist()
            xlsx_data["tables"].append(table)

        if not xlsx_data["tables"]:
            xlsx_data["text"] = "[No data found in any Excel sheets]"

    except Exception as e:
        xlsx_data["text"] = f"[Error parsing XLSX file: {e}]"

    return xlsx_data

def parse_file_ordered(filepath):
    ext = filepath.lower().split('.')[-1]
    if ext == "pdf":
        return parse_pdf_ordered(filepath)
    elif ext == "docx":
        return parse_docx_ordered(filepath)
    elif ext == "txt":
        text_data = parse_txt(filepath)
        return [{"type": "text", "content": text_data.get("text", "")}]
    elif ext == "csv":
        csv_data = parse_csv(filepath)
        # treat CSV as tables
        tables = csv_data.get("tables", [])
        content = []
        for table in tables:
            content.append({"type": "table", "content": table})
        return content
    elif ext == "xlsx":
        xlsx_data = parse_xlsx(filepath)
        content = []
        for table in xlsx_data.get("tables", []):
            content.append({"type": "table", "content": table})
        text = xlsx_data.get("text", "")
        if text:
            content.append({"type": "text", "content": text})
        return content
    elif ext in ['rtf', 'doc', 'wpd', 'wps']:
        converted = convert_with_libreoffice(filepath)
        if converted and os.path.exists(converted):
            return parse_docx_ordered(converted)
        else:
            text_data = parse_with_textract(filepath)
            return [{"type": "text", "content": text_data.get("text", "")}]
    else:
        return [{"type": "text", "content": f"Unsupported file type: {ext}"}]


def format_jsonl_ordered(content_list):
    jsonl_entries = []
    for item in content_list:
        entry = {
            "type": item.get("type"),
            "content": item.get("content")
        }
        # Optional metadata
        if "line_number" in item:
            entry["line_number"] = item["line_number"]
        if "page" in item:
            entry["page"] = item["page"]
        if "y0" in item:
            entry["y0"] = item["y0"]
        jsonl_entries.append(entry)
    return jsonl_entries

def write_jsonl_from_formatted(content_list, output_path):
    dir_path = os.path.dirname(output_path)
    if dir_path:
        os.makedirs(dir_path, exist_ok=True)
    with open(output_path, 'w', encoding='utf-8') as f:
        for item in content_list:
            json.dump(item, f, ensure_ascii=False)
            f.write('\n')

def run_parser(filepath):
    result = parse_file_ordered(filepath)
    output_path = os.path.splitext(filepath)[0] + "_output.jsonl"
    write_jsonl_from_formatted(result, output_path)
    return output_path  # ✅ return path for programmatic use


if __name__ == "__main__":
    import sys
    if len(sys.argv) != 2:
        print("Usage: python universal_parser.py <file_path>")
        sys.exit(1)

    filepath = sys.argv[1]
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        sys.exit(1)

    result = run_parser(filepath)

    output_path = os.path.splitext(filepath)[0] + "_output.jsonl"

    # Write JSONL output
    write_jsonl_from_formatted(result, output_path)

    print(f"✅ JSONL output saved to: {output_path}")
